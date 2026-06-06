import os
import tempfile
import logging
import traceback
import subprocess

class DocxPreviewService:
    def __init__(self):
        self._temp_files = []

    def create_preview_pdf(self, docx_path):
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            
        fd, temp_pdf_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        
        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(temp_pdf_path)
        
        logging.info(f"Starting DOCX conversion. Input: {abs_docx}, Target Output: {abs_pdf}")
        
        def is_word_installed():
            if os.name != 'nt': return False
            try:
                import winreg
                key = winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, "Word.Application")
                winreg.CloseKey(key)
                return True
            except Exception:
                return False

        word_installed = is_word_installed()
        logging.info(f"Microsoft Word installed: {word_installed}")
        
        conversion_success = False
        
        if word_installed:
            try:
                from docx2pdf import convert
                import sys
                import io
                from contextlib import redirect_stdout, redirect_stderr
                
                logging.info("Attempting docx2pdf conversion...")
                logging.info(f"docx2pdf: Redirecting sys.stdout/stderr to prevent tqdm crash in console=False mode.")
                
                # Protect tqdm/docx2pdf from crashing when sys.stdout is None (PyInstaller windowed mode)
                dummy_out = io.StringIO()
                dummy_err = io.StringIO()
                with redirect_stdout(dummy_out), redirect_stderr(dummy_err):
                    convert(abs_docx, abs_pdf)
                    
                logging.info("docx2pdf conversion executed successfully.")
                conversion_success = True
            except Exception as e:
                err_tb = traceback.format_exc()
                logging.error(f"docx2pdf conversion failed:\n{err_tb}")
                
        lo_exe = None
        if not conversion_success:
            logging.info("Attempting LibreOffice headless conversion fallback...")
            lo_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                "soffice"
            ]
            for p in lo_paths:
                import shutil
                if shutil.which(p) or os.path.exists(p):
                    lo_exe = p if os.path.exists(p) else shutil.which(p)
                    break
                    
            if lo_exe:
                try:
                    out_dir = os.path.dirname(abs_pdf)
                    cmd = [lo_exe, "--headless", "--convert-to", "pdf", abs_docx, "--outdir", out_dir]
                    result = subprocess.run(cmd, check=True, capture_output=True, text=True)
                    logging.info(f"LibreOffice output: {result.stdout}")
                    
                    base_name = os.path.splitext(os.path.basename(abs_docx))[0]
                    lo_generated_pdf = os.path.join(out_dir, base_name + ".pdf")
                    
                    if os.path.exists(lo_generated_pdf):
                        import shutil
                        if os.path.abspath(lo_generated_pdf) != os.path.abspath(abs_pdf):
                            if os.path.exists(abs_pdf):
                                os.remove(abs_pdf)
                            shutil.move(lo_generated_pdf, abs_pdf)
                        conversion_success = True
                    else:
                        logging.error("LibreOffice command succeeded but expected output PDF not found.")
                except subprocess.CalledProcessError as e:
                    logging.error(f"LibreOffice process failed with code {e.returncode}.\nSTDOUT: {e.stdout}\nSTDERR: {e.stderr}")
                except Exception as fallback_e:
                    fallback_tb = traceback.format_exc()
                    logging.error(f"LibreOffice conversion failed:\n{fallback_tb}")
            else:
                logging.info("LibreOffice soffice executable not found.")
                
        if not conversion_success:
            logging.info("Both docx2pdf and LibreOffice unavailable/failed. Using python-docx fallback rendering.")
            try:
                self._generate_fallback_pdf(abs_docx, abs_pdf)
                conversion_success = True
            except Exception as e:
                fallback_tb = traceback.format_exc()
                logging.error(f"python-docx fallback rendering failed:\n{fallback_tb}")
                
        if not conversion_success:
            if os.path.exists(temp_pdf_path):
                try: os.remove(temp_pdf_path)
                except: pass
            
            if not word_installed and not lo_exe:
                raise RuntimeError("Cannot preview DOCX: Microsoft Word and LibreOffice are not installed, and basic fallback rendering failed.")
            else:
                raise RuntimeError("All DOCX preview conversion methods failed. Please check the logs for detailed errors.")
            
        if not os.path.exists(abs_pdf):
            raise FileNotFoundError("PDF file was not created.")
            
        file_size = os.path.getsize(abs_pdf)
        logging.info(f"Conversion finished. Output PDF: {abs_pdf}, Size: {file_size} bytes")
        
        if file_size == 0:
            if os.path.exists(abs_pdf):
                try: os.remove(abs_pdf)
                except: pass
            raise ValueError("PDF file size is 0 bytes.")
            
        self._temp_files.append(abs_pdf)
        return abs_pdf

    def _generate_fallback_pdf(self, docx_path, pdf_path):
        import docx
        from PyQt6.QtGui import QTextDocument, QPdfWriter, QPageSize, QPageLayout
        from PyQt6.QtCore import QMarginsF
        
        doc = docx.Document(docx_path)
        
        html = "<html><body style='font-family: sans-serif;'>"
        html += "<div style='color: #666; border-bottom: 1px solid #ccc; padding-bottom: 10px; margin-bottom: 20px;'>"
        html += "<h2>Basic Preview Mode</h2>"
        html += "<p>Microsoft Word or LibreOffice could not be found to render a high-fidelity preview.</p>"
        html += "</div>"
        
        for p in doc.paragraphs:
            if p.text.strip():
                html += f"<p>{p.text}</p>"
            else:
                html += "<br>"
                
        for table in doc.tables:
            html += "<table border='1' style='border-collapse: collapse; width: 100%;'>"
            for row in table.rows:
                html += "<tr>"
                for cell in row.cells:
                    html += f"<td style='padding: 5px;'>{cell.text}</td>"
                html += "</tr>"
            html += "</table><br>"
            
        html += "</body></html>"
        
        writer = QPdfWriter(pdf_path)
        writer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
        layout = QPageLayout()
        layout.setMargins(QMarginsF(20, 20, 20, 20))
        writer.setPageLayout(layout)
        
        text_doc = QTextDocument()
        text_doc.setHtml(html)
        text_doc.print(writer)

    def cleanup(self):
        for temp_file in self._temp_files:
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except Exception:
                    pass
        self._temp_files.clear()
        
    def __del__(self):
        self.cleanup()
