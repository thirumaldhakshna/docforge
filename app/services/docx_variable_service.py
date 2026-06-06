from docx import Document
import re

class DocxVariableService:
    def __init__(self):
        self.pattern_curly = re.compile(r'\{\{([^}]+)\}\}')
        self.pattern_dollar = re.compile(r'\$([a-zA-Z0-9_]+)')

    def extract_variables(self, doc_path):
        doc = Document(doc_path)
        import logging
        logging.info(f"DocxVariableService: Loaded DOCX for variable extraction from {doc_path}")
        all_vars = set()
        
        def scan_text(text):
            if not text: return
            vars_curly = self.pattern_curly.findall(text)
            vars_dollar = self.pattern_dollar.findall(text)
            for v in vars_curly + vars_dollar:
                all_vars.add(v.strip())

        # Scan standalone paragraphs
        for para in doc.paragraphs:
            scan_text(para.text)

        # Scan tables recursively
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        scan_text(para.text)

        result = sorted(list(all_vars))
        logging.info(f"DocxVariableService: Extracted {len(result)} variables: {result}")
        return result
